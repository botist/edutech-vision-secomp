from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from xml.sax.saxutils import escape


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "entrega_professor"
DOCX_OUTPUT = OUTPUT_DIR / "artigo_edutech_vision_grupo3.docx"
PDF_OUTPUT = OUTPUT_DIR / "artigo_edutech_vision_grupo3.pdf"
FIGURE_OUTPUT = ROOT / "tmp" / "artigo_sbc" / "pipeline_generated.png"

TITLE = "EduTech Vision: Atenção, Postura e Engajamento por Visão Computacional em Tempo Real"
AUTHORS = "Fernando Cobianchi, Samuel Bernini"
INSTITUTION_LINES = [
    "Bacharelado em Ciência da Computação -- Universidade do Estado de Mato Grosso (UNEMAT)",
    "Campus Universitário de Rondonópolis -- Rondonópolis -- MT -- Brasil",
    "{fernando.cobianchi, samuel.bernini}@unemat.br",
]


ABSTRACT = (
    "EduTech Vision is a real-time digital image processing system for educational demonstrations "
    "of visual attention, posture and audience engagement. The standard pipeline combines OpenCV "
    "YuNet face detection, expanded face crops, MediaPipe Face Landmarker, EAR/MAR geometric "
    "metrics, solvePnP head pose, moving averages and temporal state machines. In a reproducible "
    "smoke benchmark at 960x540, the enhanced profile improved face recall from 0.193 to 0.592 "
    "in individual mode and from 0.011 to 0.443 in audience mode, with F1 of 0.712 and 0.583. "
    "A specific individual-mode audit processed 15,193 full-video frames and obtained 86.6% valid "
    "face/metric frames and 82.5% frontal frames. The system is a computer vision demonstration, "
    "not a clinical or pedagogical diagnostic tool."
)

RESUMO = (
    "Este artigo apresenta o EduTech Vision, um sistema de processamento digital de imagens em "
    "tempo real para demonstrações educacionais de atenção visual, postura e engajamento de "
    "plateia. O pipeline padrão combina detecção facial OpenCV YuNet, recortes faciais ampliados, "
    "MediaPipe Face Landmarker, métricas geométricas EAR/MAR, pose de cabeça por solvePnP, médias "
    "móveis e máquinas de estados temporais. Em benchmark smoke reprodutível a 960x540, o perfil "
    "enhanced elevou o recall de face de 0,193 para 0,592 no modo individual e de 0,011 para "
    "0,443 no modo plateia, com F1 de 0,712 e 0,583. Uma auditoria específica do modo individual "
    "processou 15.193 frames de vídeos completos e obteve 86,6% de frames com face/métricas "
    "válidas e 82,5% de frames frontais. O sistema é uma demonstração de visão computacional, "
    "não uma ferramenta diagnóstica clínica ou pedagógica."
)


@dataclass(frozen=True)
class Section:
    title: str
    paragraphs: tuple[str, ...]


SECTIONS = [
    Section(
        "Introdução",
        (
            "A observação de atenção, postura e engajamento em contextos educacionais costuma "
            "depender da percepção subjetiva do professor ou do próprio estudante. Essa percepção "
            "pode ser pedagogicamente útil, mas é descontínua, difícil de registrar e pouco "
            "reprodutível. Em Processamento Digital de Imagens (PDI), sinais visuais como abertura "
            "dos olhos, abertura da boca, orientação da cabeça e quantidade de faces no quadro "
            "permitem demonstrar como imagens podem ser transformadas em indicadores geométricos "
            "e temporais.",
            "O EduTech Vision foi desenvolvido na Trilha 3 da disciplina de PDI com dois cenários "
            "complementares. No modo individual, uma câmera frontal acompanha uma pessoa e estima "
            "fadiga ocular, bocejo, queda postural e desatenção sustentada. No modo plateia, uma "
            "câmera voltada a um grupo estima, de forma agregada, a proporção de faces com pose "
            "orientada ao palco ou professor. O sistema não realiza reconhecimento de identidade "
            "e não mede aprendizagem.",
            "O objetivo deste artigo é apresentar a versão final do produto, sua fundamentação "
            "técnica, o pipeline de processamento, a metodologia experimental e os resultados "
            "quantitativos disponíveis. A escrita prioriza evidências numéricas e limitações, pois "
            "os indicadores gerados são aproximações visuais para demonstração acadêmica de PDI "
            "em tempo real.",
        ),
    ),
    Section(
        "Fundamentação Teórica",
        (
            "PDI e visão computacional organizam etapas de aquisição, pré-processamento, "
            "segmentação, extração de características e interpretação visual [Gonzalez and Woods "
            "2018, Szeliski 2022]. Em aplicações com faces, detectores localizam regiões de "
            "interesse, enquanto modelos de landmarks estimam pontos faciais que permitem calcular "
            "métricas geométricas e pose.",
            "No modo individual, o Eye Aspect Ratio (EAR) é calculado pela razão entre distâncias "
            "verticais e horizontais de landmarks dos olhos, seguindo a ideia de detecção de "
            "piscadas por geometria facial [Soukupova and Cech 2016]. Quando os olhos fecham, as "
            "distâncias verticais diminuem e o EAR cai. A Mouth Aspect Ratio (MAR) usa princípio "
            "semelhante para representar abertura da boca e sinalizar bocejos sustentados.",
            "A orientação da cabeça é estimada por correspondências 2D-3D e pela família de "
            "algoritmos solvePnP [Lepetit et al. 2009]. A partir de landmarks como nariz, queixo, "
            "cantos dos olhos e boca, o sistema estima yaw, pitch e roll. Esses ângulos são "
            "suavizados no tempo porque frames isolados podem conter ruído de landmark, oclusão "
            "ou blur de movimento.",
            "A detecção inicial de faces foi avaliada contra o uso direto de MediaPipe no quadro "
            "completo [Lugaresi et al. 2019]. O benchmark WIDER FACE [Yang et al. 2016] foi usado "
            "como base anotada para medir recall, precisão e F1 em diferentes tamanhos de face, "
            "distâncias e perturbações sintéticas.",
        ),
    ),
    Section(
        "Metodologia e Pipeline",
        (
            "A Figura 1 resume o pipeline operacional. A entrada é uma webcam ou arquivo de vídeo "
            "redimensionado para 960x540 no modo padrão. O perfil enhanced aplica OpenCV YuNet "
            "para localizar faces, usa supressão de não máximos e recortes ampliados, e então "
            "executa MediaPipe Face Landmarker em cada região de interesse. Essa escolha preserva "
            "o MediaPipe para landmarks e desloca o gargalo de faces pequenas para um detector "
            "frontal mais adequado.",
            "Três perfis foram mantidos para comparação metodológica. O perfil mediapipe é o "
            "baseline, com processamento direto do quadro. O perfil enhanced é o padrão "
            "distribuível, combinando YuNet e landmarks em crops. O perfil research usa "
            "SCRFD/InsightFace apenas como comparativo acadêmico opcional; o padrão de entrega "
            "é o perfil enhanced.",
            "No modo individual, o sistema usa no máximo uma face e confiança facial 0,70. A "
            "calibração dura 5 s a partir da primeira amostra facial válida, evitando baseline "
            "vazio quando a câmera ou vídeo inicia sem rosto. O baseline armazena medianas e "
            "desvios robustos de EAR, MAR, yaw e pitch. Para yaw/pitch, a estatística é circular: "
            "valores próximos de +179 e -179 graus são tratados como orientações equivalentes, "
            "não como extremos opostos.",
            "Os alertas usam médias móveis e condições sustentadas. Fadiga ocular exige EAR médio "
            "abaixo do limiar dinâmico e queda simultânea dos dois olhos; bocejo exige MAR acima "
            "do limiar; postura usa desvio de pitch; e desatenção usa desvio de yaw. No modo "
            "plateia, cada face com pose válida é comparada a um eixo de palco calibrado, e os "
            "resultados são agregados em janelas de 10 s.",
        ),
    ),
    Section(
        "Metodologia Experimental",
        (
            "A avaliação foi organizada em quatro frentes. A primeira foi um benchmark smoke do "
            "detector, comparando o baseline mediapipe e o perfil enhanced nos modos individual "
            "e plateia. As métricas foram recall de face, recall de landmarks, precisão, F1, FPS "
            "médio e FPS mínimo. O recall de face mede caixas detectadas; o recall de landmarks "
            "exige também pontos suficientes para EAR, MAR ou pose.",
            "A segunda frente foi uma auditoria específica do modo individual. Foram baixados "
            "vídeos frontais 1080p e retratos estáticos por manifesto reprodutível. O auditor "
            "processou vídeos completos, sem amostragem por intervalo, redimensionando cada frame "
            "para a mesma resolução operacional do app. Cada frame gerou uma linha com detecção, "
            "qualidade dos landmarks, EAR, MAR, yaw, pitch, deltas em relação ao baseline e "
            "classificação frontal/neutra.",
            "A terceira frente foi a auditoria HQ de plateia. Vídeos 1080p completos foram "
            "processados do primeiro ao último frame e quadros selecionados foram comparados a "
            "contagens manuais de rostos avaliáveis. Foram medidos acertos exatos, erro de no "
            "máximo um rosto, erro absoluto médio, presença da contagem manual em janela temporal "
            "de +/-1 s, FPS médio e FPS p10.",
            "A quarta frente foi a instrumentação dos protocolos formais: iluminação, oclusão, "
            "FPS, matriz de confusão e tolerância a falhas. Na versão atual, apenas a matriz "
            "demonstrativa possui resultado consolidado; os demais protocolos estão implementados "
            "como scripts de coleta, mas não foram usados como evidência final por falta de nova "
            "execução presencial completa. Essa limitação é explicitada para evitar uso indevido "
            "de logs preliminares.",
        ),
    ),
]

TABLE_BENCHMARK = [
    ["Modo", "Perfil", "Recall face", "Recall landmarks", "Precisão", "F1", "FPS médio"],
    ["Individual", "mediapipe", "0,193", "0,193", "1,000", "0,324", "242,4"],
    ["Individual", "enhanced", "0,592", "0,294", "0,894", "0,712", "27,6"],
    ["Plateia", "mediapipe", "0,011", "0,011", "1,000", "0,022", "491,6"],
    ["Plateia", "enhanced", "0,443", "0,089", "0,851", "0,583", "74,5"],
]

TABLE_INDIVIDUAL_AUDIT = [
    ["Métrica", "Valor"],
    ["Vídeos processados", "3"],
    ["Frames totais", "15.193"],
    ["Frames com face/métricas válidas", "86,6%"],
    ["Frames frontais", "82,5%"],
    ["Frames neutros em pose", "76,7%"],
    ["Frames aprovados", "73,0%"],
    ["Yaw delta mediano", "4,0 graus"],
    ["Pitch delta mediano", "1,8 grau"],
]

TABLE_AUDIENCE_AUDIT = [
    ["Configuração", "Exatos", "Erro <=1", "Erro médio", "Manual em +/-1s", "FPS médio", "FPS p10"],
    ["enhanced c0,60 m24", "5/11", "10/11", "0,636", "10/11", "13,3", "9,4"],
]

RESULT_PARAGRAPHS = [
    "A Tabela 1 apresenta o benchmark smoke canônico da versão atual. Em ambos os modos, o "
    "perfil enhanced aumenta substancialmente o recall de face e o F1 em relação ao baseline "
    "MediaPipe, mantendo desempenho em tempo real. A redução de FPS é esperada porque o pipeline "
    "passa a executar detecção por YuNet, recortes ampliados e landmarks em regiões de interesse.",
    "A Tabela 2 mostra a auditoria específica do modo individual. Nos três vídeos frontais "
    "completos, o sistema processou 15.193 frames. Os frames sem face dos vídeos NASA correspondem "
    "principalmente a cartelas e b-roll de abertura, não a perda sistemática de rosto frontal. "
    "Quando havia pessoa frontal no quadro, os deltas medianos foram pequenos: 4,0 graus em yaw "
    "e 1,8 grau em pitch.",
    "A Tabela 3 resume a auditoria HQ de plateia com a configuração enhanced, confiança 0,60 e "
    "limite de 24 faces. Essa avaliação é mais pesada que o benchmark smoke porque usa vídeos "
    "1080p completos e compara a saída temporal do app com contagens manuais.",
    "A matriz de confusão disponível contém 50 amostras sintéticas de atenção/desatenção, usadas "
    "para testar o fluxo de cálculo, visualização e relatório. O resultado foi acurácia, precisão "
    "macro, recall macro e F1 macro de 0,92. Por serem amostras demonstrativas, esses valores não "
    "substituem uma validação científica com rotulação real.",
]

DISCUSSION = Section(
    "Discussão",
    (
        "Os resultados justificam o perfil enhanced como padrão. O baseline mediapipe é muito "
        "rápido, mas perde faces pequenas ou distantes, sobretudo no modo plateia. A combinação "
        "YuNet + crops + landmarks aumenta a utilidade visual do produto sem exigir treinamento "
        "próprio. Separar recall de face e recall de landmarks evita uma interpretação incorreta: "
        "uma caixa pode estar correta, mas ainda não oferecer pontos confiáveis para pose ou EAR/MAR.",
        "No modo individual, os principais ajustes metodológicos foram motivados pelos testes. O "
        "pitch retornado por solvePnP pode alternar entre +179 e -179 graus para faces frontais; "
        "por isso o baseline passou a usar média circular. A calibração também passou a esperar "
        "a primeira face válida, pois vídeos com abertura sem rosto geravam baseline inadequado. "
        "Esses dois pontos aumentaram a consistência dos deltas de pose em vídeos frontais.",
        "Ainda há limitações importantes. Iluminação irregular, contraluz, oclusões, rosto em "
        "perfil, distância excessiva, motion blur e cortes nas bordas reduzem a estabilidade de "
        "landmarks. No modo plateia, a estimativa é agregada e visual; ela não mede atenção "
        "cognitiva. No modo individual, EAR e MAR variam entre pessoas e dependem de calibração "
        "inicial.",
    ),
)

ETHICS = Section(
    "Considerações Éticas e LGPD",
    (
        "Como o sistema captura imagem de pessoas, seu uso deve seguir minimização de dados, "
        "transparência e finalidade acadêmica. O modo individual deve ser demonstrado com "
        "consentimento do participante e sem finalidade punitiva, clínica ou classificatória. O "
        "modo plateia deve operar com métricas agregadas e não deve identificar alunos, associar "
        "nomes a rostos ou produzir avaliação individual de aprendizagem.",
        "A proposta privilegia processamento local e armazenamento apenas de telemetria numérica, "
        "como séries temporais, contagens, percentuais agregados e eventos de interface. Em "
        "demonstrações públicas, recomenda-se sinalização visível, opção de ficar fora do campo "
        "da câmera e uso de voluntários quando forem necessárias evidências visuais.",
    ),
)

CONCLUSION = Section(
    "Conclusão",
    (
        "Este artigo apresentou o EduTech Vision como aplicação de PDI em tempo real para "
        "monitoramento aproximado de atenção visual, postura e engajamento agregado. A arquitetura "
        "final combina YuNet, MediaPipe Face Landmarker em recortes ampliados, EAR, MAR, solvePnP, "
        "suavização temporal, painel OpenCV e relatórios de sessão.",
        "As evidências quantitativas mostram ganho claro do perfil enhanced sobre o baseline "
        "mediapipe e uma auditoria específica mais forte para o modo individual. Ao mesmo tempo, "
        "o artigo delimita o escopo: o sistema é uma demonstração acadêmica e não uma ferramenta "
        "diagnóstica ou avaliativa. Trabalhos futuros incluem reexecutar os cinco protocolos "
        "formais com coleta presencial completa, ampliar amostras reais rotuladas e testar "
        "condições controladas de iluminação, distância e oclusão.",
    ),
)

ACKNOWLEDGMENTS = (
    "Os autores agradecem ao curso de Bacharelado em Ciência da Computação da UNEMAT e ao docente "
    "da disciplina de Processamento Digital de Imagens pelas diretrizes, críticas e critérios de "
    "avaliação do projeto. Ferramentas de IA generativa foram utilizadas pontualmente como apoio "
    "à revisão gramatical e à organização textual; o conteúdo técnico, a interpretação dos "
    "resultados e a responsabilidade final permanecem dos autores."
)

REFERENCES = [
    "Brasil (2018). Lei nº 13.709, de 14 de agosto de 2018. Lei Geral de Proteção de Dados Pessoais (LGPD).",
    "Gonzalez, R. C. and Woods, R. E. (2018). Digital Image Processing. 4. ed. Pearson.",
    "Lepetit, V., Moreno-Noguer, F. and Fua, P. (2009). EPnP: An Accurate O(n) Solution to the PnP Problem. International Journal of Computer Vision, 81(2):155-166.",
    "Lugaresi, C. et al. (2019). MediaPipe: A Framework for Building Perception Pipelines. arXiv:1906.08172.",
    "OpenCV Zoo (2026). YuNet Face Detection Model. Disponível em: https://github.com/opencv/opencv_zoo.",
    "Soukupova, T. and Cech, J. (2016). Real-Time Eye Blink Detection using Facial Landmarks. 21st Computer Vision Winter Workshop.",
    "Szeliski, R. (2022). Computer Vision: Algorithms and Applications. 2. ed. Springer.",
    "Yang, S., Luo, P., Loy, C. C. and Tang, X. (2016). WIDER FACE: A Face Detection Benchmark. Proceedings of CVPR.",
]


def set_cell_text(cell, text: str, bold: bool = False, size: int = 8) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)


def set_paragraph_font(paragraph, *, size: int = 12, bold: bool = False, italic: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic


def add_centered(document: Document, text: str, *, bold: bool = False, size: int = 12) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    set_paragraph_font(paragraph, size=size, bold=bold)


def add_abstract(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Cm(0.8)
    paragraph.paragraph_format.right_indent = Cm(0.8)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_after = Pt(8)
    label_run = paragraph.add_run(f"{label}. ")
    label_run.bold = True
    label_run.italic = True
    body_run = paragraph.add_run(text)
    body_run.italic = True
    set_paragraph_font(paragraph, italic=True)


def add_section_heading(document: Document, title: str) -> None:
    paragraph = document.add_paragraph(title)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    set_paragraph_font(paragraph, size=13, bold=True)


def add_body_paragraph(document: Document, text: str, *, first: bool = False) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0 if first else 1.27)
    paragraph.paragraph_format.space_after = Pt(6)
    set_paragraph_font(paragraph)


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(7)
    set_paragraph_font(paragraph, size=10, bold=True)


def add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for cell, value in zip(table.rows[0].cells, rows[0]):
        set_cell_text(cell, value, bold=True)
    for row in rows[1:]:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell_text(cell, value)
    document.add_paragraph().paragraph_format.space_after = Pt(4)


def draw_pipeline_figure() -> None:
    FIGURE_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    labels = [
        "Entrada\nwebcam/vídeo",
        "YuNet\n+ NMS",
        "Crop facial\nampliado",
        "MediaPipe\nLandmarker",
        "EAR/MAR\n+ solvePnP",
        "Médias móveis\n+ estados",
        "Alertas\n+ agregação 10 s",
    ]
    fig, ax = plt.subplots(figsize=(12, 2.7), dpi=180)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    x_positions = [0.05, 0.19, 0.33, 0.47, 0.61, 0.75, 0.89]
    width = 0.115
    height = 0.42
    for idx, (x, label) in enumerate(zip(x_positions, labels)):
        rect = plt.Rectangle(
            (x - width / 2, 0.42),
            width,
            height,
            facecolor="#f7f7f7",
            edgecolor="#222222",
            linewidth=1.1,
        )
        ax.add_patch(rect)
        ax.text(x, 0.63, label, ha="center", va="center", fontsize=8.5, color="#111111")
        if idx < len(x_positions) - 1:
            ax.annotate(
                "",
                xy=(x_positions[idx + 1] - width / 2 - 0.01, 0.63),
                xytext=(x + width / 2 + 0.01, 0.63),
                arrowprops={"arrowstyle": "->", "lw": 1.0, "color": "#222222"},
            )
    ax.text(
        0.5,
        0.18,
        "Saídas: painel OpenCV, logs CSV, gráficos, relatório HTML/PDF e métricas agregadas de plateia",
        ha="center",
        va="center",
        fontsize=9,
        color="#222222",
    )
    fig.tight_layout(pad=0.2)
    fig.savefig(FIGURE_OUTPUT, bbox_inches="tight")
    plt.close(fig)


def build_docx() -> None:
    draw_pipeline_figure()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_document_defaults(document)

    add_centered(document, TITLE, bold=True, size=16)
    add_centered(document, AUTHORS, bold=True)
    for line in INSTITUTION_LINES:
        add_centered(document, line)
    add_abstract(document, "Abstract", ABSTRACT)
    add_abstract(document, "Resumo", RESUMO)

    for section in SECTIONS:
        add_section_heading(document, section.title)
        for idx, paragraph in enumerate(section.paragraphs):
            add_body_paragraph(document, paragraph, first=idx == 0)
        if section.title == "Metodologia e Pipeline":
            picture = document.add_picture(str(FIGURE_OUTPUT), width=Cm(14.5))
            picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_caption(document, "Figura 1. Pipeline atual do EduTech Vision.")

    add_section_heading(document, "Resultados")
    add_body_paragraph(document, RESULT_PARAGRAPHS[0], first=True)
    add_caption(document, "Tabela 1. Benchmark smoke do detector em resolução 960x540.")
    add_table(document, TABLE_BENCHMARK)
    add_body_paragraph(document, RESULT_PARAGRAPHS[1], first=False)
    add_caption(document, "Tabela 2. Auditoria do modo individual em vídeos frontais completos.")
    add_table(document, TABLE_INDIVIDUAL_AUDIT)
    add_body_paragraph(document, RESULT_PARAGRAPHS[2], first=False)
    add_caption(document, "Tabela 3. Auditoria HQ de plateia com enhanced, confiança 0,60 e 24 faces.")
    add_table(document, TABLE_AUDIENCE_AUDIT)
    add_body_paragraph(document, RESULT_PARAGRAPHS[3], first=False)

    for section in [DISCUSSION, ETHICS, CONCLUSION]:
        add_section_heading(document, section.title)
        for idx, paragraph in enumerate(section.paragraphs):
            add_body_paragraph(document, paragraph, first=idx == 0)

    add_section_heading(document, "Agradecimentos")
    add_body_paragraph(document, ACKNOWLEDGMENTS, first=True)
    add_section_heading(document, "Referências")
    for reference in REFERENCES:
        add_body_paragraph(document, reference, first=True)

    document.save(DOCX_OUTPUT)


def pdf_paragraph(text: str, style: ParagraphStyle) -> Paragraph:
    safe = escape(text)
    safe = safe.replace("Abstract.", "<b><i>Abstract.</i></b>", 1)
    safe = safe.replace("Resumo.", "<b><i>Resumo.</i></b>", 1)
    return Paragraph(safe, style)


def pdf_table(rows: list[list[str]], styles) -> Table:
    ncols = len(rows[0])
    page_width = A4[0] - 6 * cm
    if ncols >= 7:
        col_widths = [2.2 * cm, 1.65 * cm, 1.55 * cm, 1.75 * cm, 1.3 * cm, 1.25 * cm, 1.25 * cm]
    elif ncols == 2:
        col_widths = [5.0 * cm, page_width - 5.0 * cm]
    else:
        col_widths = [page_width / ncols] * ncols
    cell_style = ParagraphStyle(
        "Cell",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=7.6,
        leading=8.8,
        alignment=TA_CENTER,
    )
    data = [[pdf_paragraph(cell, cell_style) for cell in row] for row in rows]
    table = Table(data, colWidths=col_widths, repeatRows=1, hAlign="CENTER")
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 2),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]
        )
    )
    return table


def build_pdf() -> None:
    draw_pipeline_figure()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleSBC",
        parent=styles["Normal"],
        fontName="Times-Bold",
        fontSize=16,
        leading=18,
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    center_bold = ParagraphStyle(
        "CenterBold",
        parent=styles["Normal"],
        fontName="Times-Bold",
        fontSize=12,
        leading=14,
        alignment=TA_CENTER,
        spaceAfter=2,
    )
    center = ParagraphStyle(
        "Center",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=12,
        leading=14,
        alignment=TA_CENTER,
        spaceAfter=2,
    )
    abstract_style = ParagraphStyle(
        "Abstract",
        parent=styles["Normal"],
        fontName="Times-Italic",
        fontSize=12,
        leading=14,
        leftIndent=0.8 * cm,
        rightIndent=0.8 * cm,
        alignment=TA_JUSTIFY,
        spaceAfter=8,
    )
    section_style = ParagraphStyle(
        "SectionSBC",
        parent=styles["Normal"],
        fontName="Times-Bold",
        fontSize=13,
        leading=15,
        alignment=TA_LEFT,
        spaceBefore=10,
        spaceAfter=6,
        keepWithNext=1,
    )
    body_first = ParagraphStyle(
        "BodyFirst",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=12,
        leading=14,
        alignment=TA_JUSTIFY,
        firstLineIndent=0,
        spaceAfter=6,
    )
    body = ParagraphStyle("Body", parent=body_first, firstLineIndent=1.27 * cm)
    caption = ParagraphStyle(
        "Caption",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=12,
        alignment=TA_CENTER,
        spaceBefore=4,
        spaceAfter=7,
    )
    reference_style = ParagraphStyle(
        "Reference",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=11,
        leading=13,
        alignment=TA_LEFT,
        spaceAfter=4,
    )

    story = [
        pdf_paragraph(TITLE, title_style),
        pdf_paragraph(AUTHORS, center_bold),
    ]
    story.extend(pdf_paragraph(line, center) for line in INSTITUTION_LINES)
    story.append(pdf_paragraph(f"Abstract. {ABSTRACT}", abstract_style))
    story.append(pdf_paragraph(f"Resumo. {RESUMO}", abstract_style))

    for section in SECTIONS:
        story.append(pdf_paragraph(section.title, section_style))
        for idx, paragraph in enumerate(section.paragraphs):
            story.append(pdf_paragraph(paragraph, body_first if idx == 0 else body))
        if section.title == "Metodologia e Pipeline":
            story.append(Image(str(FIGURE_OUTPUT), width=14.5 * cm, height=3.3 * cm))
            story.append(pdf_paragraph("Figura 1. Pipeline atual do EduTech Vision.", caption))

    story.append(pdf_paragraph("Resultados", section_style))
    story.append(pdf_paragraph(RESULT_PARAGRAPHS[0], body_first))
    story.append(pdf_paragraph("Tabela 1. Benchmark smoke do detector em resolução 960x540.", caption))
    story.append(pdf_table(TABLE_BENCHMARK, styles))
    story.append(Spacer(1, 7))
    story.append(pdf_paragraph(RESULT_PARAGRAPHS[1], body))
    story.append(pdf_paragraph("Tabela 2. Auditoria do modo individual em vídeos frontais completos.", caption))
    story.append(pdf_table(TABLE_INDIVIDUAL_AUDIT, styles))
    story.append(Spacer(1, 7))
    story.append(pdf_paragraph(RESULT_PARAGRAPHS[2], body))
    story.append(pdf_paragraph("Tabela 3. Auditoria HQ de plateia com enhanced, confiança 0,60 e 24 faces.", caption))
    story.append(pdf_table(TABLE_AUDIENCE_AUDIT, styles))
    story.append(Spacer(1, 7))
    story.append(pdf_paragraph(RESULT_PARAGRAPHS[3], body))

    for section in [DISCUSSION, ETHICS, CONCLUSION]:
        story.append(pdf_paragraph(section.title, section_style))
        for idx, paragraph in enumerate(section.paragraphs):
            story.append(pdf_paragraph(paragraph, body_first if idx == 0 else body))

    story.append(pdf_paragraph("Agradecimentos", section_style))
    story.append(pdf_paragraph(ACKNOWLEDGMENTS, body_first))
    story.append(pdf_paragraph("Referências", section_style))
    for reference in REFERENCES:
        story.append(pdf_paragraph(reference, reference_style))

    document = SimpleDocTemplate(
        str(PDF_OUTPUT),
        pagesize=A4,
        rightMargin=3 * cm,
        leftMargin=3 * cm,
        topMargin=3.5 * cm,
        bottomMargin=2.5 * cm,
    )
    document.build(story)


def main() -> None:
    build_docx()
    build_pdf()
    print(f"DOCX gerado: {DOCX_OUTPUT}")
    print(f"PDF gerado: {PDF_OUTPUT}")


if __name__ == "__main__":
    main()
